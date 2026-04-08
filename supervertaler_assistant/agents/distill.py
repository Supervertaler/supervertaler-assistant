"""
supervertaler_assistant.agents.distill
======================================

Distill agent – extracts knowledge from translation-industry source
files into a single Markdown article ready for the 00_INBOX/.

Supported formats
-----------------
- **TMX** (.tmx) – translation memory exchange (XML)
- **DOCX** (.docx) – Microsoft Word documents
- **PDF** (.pdf) – PDF documents
- **TBX** (.tbx) – termbase exchange (XML)

Each format gets its own extractor that returns plain text plus a
short metadata summary. All four feed into the same LLM call using
``distill.md`` as the system prompt. The reply (a single Markdown
article) is written to ``00_INBOX/`` with a timestamped filename.

Design choices
--------------
- **Extractors degrade gracefully.** PDF and DOCX extraction can fail
  on exotic files – when they do, the agent reports the error rather
  than crashing. The user can then fall back to copy-pasting text.
- **Size cap per file.** Huge PDFs (book-length documents) are
  truncated to ~200k characters so one massive file can't blow the
  LLM context budget. The user is warned when truncation happens.
- **Optional dependencies.** ``python-docx``, ``pypdf`` and ``lxml``
  are pulled in via the ``distill`` extra in pyproject.toml.
  Importing them lazily means the chat-only code paths don't pay the
  cost of loading them at startup.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Literal

from ..llm import LlmClient, system, user

_LOG = logging.getLogger(__name__)


SourceFormat = Literal["tmx", "docx", "pdf", "tbx"]

# Per-file char cap before we start truncating. ~200k chars ≈ 50k tokens,
# generous enough for most real docs.
_MAX_EXTRACT_CHARS = 200_000


# ─── Extraction result ─────────────────────────────────────────────────────


@dataclass
class Extracted:
    """One extracted-text payload ready for LLM distillation."""

    text: str
    metadata: dict
    format: SourceFormat
    source_file: str
    truncated: bool = False


@dataclass
class DistillResult:
    """Outcome of one Distill run."""

    source_file: Path
    format: SourceFormat | None = None
    inbox_file: Path | None = None
    truncated: bool = False
    llm_raw_reply: str = ""
    error: str | None = None

    @property
    def ok(self) -> bool:
        return self.error is None and self.inbox_file is not None

    def summary(self) -> str:
        if self.error:
            return f"Distill failed: {self.error}"
        if self.inbox_file is None:
            return "Distill produced no output."

        bits = [f"Distilled {self.source_file.name} → {self.inbox_file.name}."]
        if self.truncated:
            bits.append(
                "⚠ Source file was truncated to fit the LLM context budget – "
                "only the first portion was distilled."
            )
        return " ".join(bits)


# ─── Agent ──────────────────────────────────────────────────────────────────


class DistillAgent:
    """Extract knowledge from a source file into an inbox-ready article."""

    def __init__(self, llm: LlmClient, memory_bank_dir: str | Path) -> None:
        self.llm = llm
        self.memory_bank_dir = Path(memory_bank_dir).resolve()
        self._prompt = _load_distill_prompt(self.memory_bank_dir)

    # ── Public API ──────────────────────────────────────────────────────

    @staticmethod
    def detect_format(path: Path) -> SourceFormat | None:
        """Return the SourceFormat for ``path`` based on its extension."""
        ext = path.suffix.lower()
        return {
            ".tmx": "tmx",
            ".docx": "docx",
            ".pdf": "pdf",
            ".tbx": "tbx",
        }.get(ext)

    def run(self, source_file: str | Path) -> DistillResult:
        """Distill a single source file into a new 00_INBOX/ article."""
        src = Path(source_file)
        result = DistillResult(source_file=src)

        fmt = self.detect_format(src)
        if fmt is None:
            result.error = (
                f"Unsupported file type: {src.suffix}. "
                "Supported: .tmx, .docx, .pdf, .tbx"
            )
            return result
        result.format = fmt

        if not src.is_file():
            result.error = f"File not found: {src}"
            return result

        try:
            extracted = _extract(src, fmt)
        except ImportError as exc:
            result.error = (
                f"Missing optional dependency for {fmt.upper()} support: {exc}. "
                "Install with: pip install 'supervertaler-assistant[distill]'"
            )
            return result
        except Exception as exc:  # noqa: BLE001 – agent boundary
            _LOG.exception("Extraction failed for %s", src)
            result.error = f"Extraction failed: {type(exc).__name__}: {exc}"
            return result

        if not extracted.text.strip():
            result.error = f"No extractable text found in {src.name}."
            return result

        result.truncated = extracted.truncated

        try:
            article = self._call_llm(extracted)
            result.llm_raw_reply = article

            inbox = self.memory_bank_dir / "00_INBOX"
            inbox.mkdir(parents=True, exist_ok=True)
            target = _pick_inbox_filename(inbox, src.stem)
            target.write_text(_ensure_trailing_newline(article), encoding="utf-8")
            result.inbox_file = target

        except Exception as exc:  # noqa: BLE001
            _LOG.exception("Distill LLM call failed")
            result.error = f"{type(exc).__name__}: {exc}"

        return result

    # ── Internals ───────────────────────────────────────────────────────

    def _call_llm(self, extracted: Extracted) -> str:
        meta_lines = [f"- **{k}**: {v}" for k, v in extracted.metadata.items()]
        meta_block = "\n".join(meta_lines) if meta_lines else "- (none)"

        user_prompt = (
            f"Source format: `{extracted.format}`\n"
            f"Source filename: `{extracted.source_file}`\n\n"
            "Metadata extracted from the source:\n"
            f"{meta_block}\n\n"
            "Extracted content follows. Distill it into a single Markdown "
            "article ready to drop into `00_INBOX/`, following the rules "
            "in your system prompt. Do NOT wrap the output in a code fence "
            "or a `### FILE:` block.\n\n"
            "```\n"
            f"{extracted.text}\n"
            "```"
        )
        return self.llm.complete([system(self._prompt), user(user_prompt)])


# ─── Format extractors ─────────────────────────────────────────────────────


def _extract(src: Path, fmt: SourceFormat) -> Extracted:
    """Dispatch to the right extractor for ``fmt``."""
    if fmt == "tmx":
        return _extract_tmx(src)
    if fmt == "tbx":
        return _extract_tbx(src)
    if fmt == "docx":
        return _extract_docx(src)
    if fmt == "pdf":
        return _extract_pdf(src)
    raise ValueError(f"unknown format: {fmt}")


# ── TMX ─────────────────────────────────────────────────────────────────────


def _extract_tmx(src: Path) -> Extracted:
    """Extract translation units from a TMX file.

    Returns a plain-text rendering of the form::

        [en-US]  source segment
        [nl-BE]  target segment
        ---
    """
    from lxml import etree  # lazy

    parser = etree.XMLParser(recover=True, huge_tree=True)
    tree = etree.parse(str(src), parser)
    root = tree.getroot()
    ns = _strip_namespace(root.tag)

    # Metadata from <header>
    header = root.find("header") if ns else None
    if header is None:
        # Try with namespace
        header = _find_any(root, "header")

    metadata: dict = {}
    if header is not None:
        for key in (
            "srclang",
            "creationtool",
            "creationtoolversion",
            "datatype",
            "segtype",
            "adminlang",
            "o-tmf",
        ):
            val = header.get(key)
            if val:
                metadata[key] = val

    # Extract TUs
    lines: list[str] = []
    tu_count = 0
    truncated = False
    total_chars = 0

    for tu in _iter_any(root, "tu"):
        tuvs = list(_iter_any(tu, "tuv"))
        tu_count += 1
        for tuv in tuvs:
            lang = (
                tuv.get("{http://www.w3.org/XML/1998/namespace}lang")
                or tuv.get("xml:lang")
                or tuv.get("lang")
                or "?"
            )
            seg = _find_any(tuv, "seg")
            if seg is None:
                continue
            # Flatten inline tags – we just want text
            text = "".join(seg.itertext()).strip()
            if not text:
                continue
            line = f"[{lang}]  {text}"
            total_chars += len(line) + 1
            if total_chars > _MAX_EXTRACT_CHARS:
                truncated = True
                break
            lines.append(line)
        if truncated:
            break
        lines.append("---")

    metadata["tu_count"] = str(tu_count)
    metadata["extracted_segments"] = str(len([ln for ln in lines if ln != "---"]))

    return Extracted(
        text="\n".join(lines),
        metadata=metadata,
        format="tmx",
        source_file=src.name,
        truncated=truncated,
    )


# ── TBX ─────────────────────────────────────────────────────────────────────


def _extract_tbx(src: Path) -> Extracted:
    """Extract term entries from a TBX termbase export.

    TBX is notoriously flexible; we do a best-effort walk over
    ``termEntry`` elements and pull any ``term`` text regardless of
    flavour (Basic, Default, Min).
    """
    from lxml import etree  # lazy

    parser = etree.XMLParser(recover=True, huge_tree=True)
    tree = etree.parse(str(src), parser)
    root = tree.getroot()

    lines: list[str] = []
    entry_count = 0
    term_count = 0
    truncated = False
    total_chars = 0

    for entry in _iter_any(root, "termEntry"):
        entry_count += 1
        entry_block: list[str] = []

        # LangSet → term(s) per language
        for langset in _iter_any(entry, "langSet"):
            lang = (
                langset.get("{http://www.w3.org/XML/1998/namespace}lang")
                or langset.get("xml:lang")
                or "?"
            )
            for term in _iter_any(langset, "term"):
                text = "".join(term.itertext()).strip()
                if text:
                    entry_block.append(f"[{lang}]  {text}")
                    term_count += 1

            # Pick up descriptive notes if present
            for note in _iter_any(langset, "note"):
                note_text = "".join(note.itertext()).strip()
                if note_text:
                    entry_block.append(f"[{lang}]  note: {note_text}")

        block_text = "\n".join(entry_block)
        if not block_text:
            continue

        chunk = block_text + "\n---\n"
        total_chars += len(chunk)
        if total_chars > _MAX_EXTRACT_CHARS:
            truncated = True
            break

        lines.append(block_text)
        lines.append("---")

    metadata = {
        "entry_count": str(entry_count),
        "term_count": str(term_count),
    }

    return Extracted(
        text="\n".join(lines),
        metadata=metadata,
        format="tbx",
        source_file=src.name,
        truncated=truncated,
    )


# ── DOCX ────────────────────────────────────────────────────────────────────


def _extract_docx(src: Path) -> Extracted:
    """Extract paragraphs and table cells from a DOCX file."""
    import docx  # python-docx – lazy

    doc = docx.Document(str(src))

    lines: list[str] = []
    total_chars = 0
    truncated = False

    def _append(line: str) -> bool:
        nonlocal total_chars, truncated
        if not line.strip():
            return True
        if total_chars + len(line) > _MAX_EXTRACT_CHARS:
            truncated = True
            return False
        lines.append(line)
        total_chars += len(line) + 1
        return True

    for para in doc.paragraphs:
        if not _append(para.text):
            break

    if not truncated:
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if not _append(row_text):
                    break
            if truncated:
                break

    metadata = {
        "paragraph_count": str(len(doc.paragraphs)),
        "table_count": str(len(doc.tables)),
    }

    return Extracted(
        text="\n".join(lines),
        metadata=metadata,
        format="docx",
        source_file=src.name,
        truncated=truncated,
    )


# ── PDF ─────────────────────────────────────────────────────────────────────


def _extract_pdf(src: Path) -> Extracted:
    """Extract text from a PDF, page by page."""
    import pypdf  # lazy

    reader = pypdf.PdfReader(str(src))
    page_count = len(reader.pages)

    chunks: list[str] = []
    total_chars = 0
    truncated = False

    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            _LOG.debug("PDF page %d extraction failed: %s", i, exc)
            continue
        text = _clean_pdf_text(text)
        if not text.strip():
            continue
        chunk = f"[Page {i + 1}]\n{text}\n"
        if total_chars + len(chunk) > _MAX_EXTRACT_CHARS:
            truncated = True
            break
        chunks.append(chunk)
        total_chars += len(chunk)

    metadata = {
        "page_count": str(page_count),
        "pages_extracted": str(len(chunks)),
    }
    if reader.metadata:
        for key in ("title", "author", "subject", "creator"):
            val = getattr(reader.metadata, key, None)
            if val:
                metadata[key] = str(val)

    return Extracted(
        text="\n".join(chunks),
        metadata=metadata,
        format="pdf",
        source_file=src.name,
        truncated=truncated,
    )


# ─── XML helpers (namespace-tolerant) ───────────────────────────────────────


def _strip_namespace(tag: str) -> str:
    """``'{ns}tag'`` → ``'tag'``."""
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _find_any(parent, local_name: str):
    """Like ``parent.find(name)`` but ignores XML namespaces."""
    for child in parent.iter():
        if _strip_namespace(child.tag) == local_name and child is not parent:
            return child
    return None


def _iter_any(parent, local_name: str):
    """Like ``parent.iter(name)`` but ignores XML namespaces."""
    for child in parent.iter():
        if _strip_namespace(child.tag) == local_name:
            yield child


# ─── Small helpers ──────────────────────────────────────────────────────────


def _clean_pdf_text(text: str) -> str:
    """Remove obvious PDF extraction artefacts (stray line numbers, form feeds)."""
    text = text.replace("\f", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _pick_inbox_filename(inbox: Path, stem: str) -> Path:
    """Return a unique path in ``inbox`` for an article derived from ``stem``.

    Filenames follow the pattern ``distilled_{stem}_{YYYYMMDD-HHMMSS}.md``
    so repeat runs never collide and the timestamp is visible in Obsidian.
    """
    safe_stem = re.sub(r"[^A-Za-z0-9 \-_().&]", "_", stem).strip() or "distilled"
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    name = f"distilled_{safe_stem}_{stamp}.md"
    return inbox / name


def _ensure_trailing_newline(text: str) -> str:
    return text if text.endswith("\n") else text + "\n"


def _load_distill_prompt(memory_bank_dir: Path) -> str:
    """Load distill.md, preferring the memory bank's copy over the bundled fallback.

    Unlike compile/lint/query, most users won't have a distill.md in
    their 06_TEMPLATES/ because it's a newer addition. That's fine –
    we always ship a bundled copy.
    """
    user_template = memory_bank_dir / "06_TEMPLATES" / "distill.md"
    if user_template.is_file():
        try:
            return user_template.read_text(encoding="utf-8")
        except OSError:
            pass

    bundled = Path(__file__).resolve().parent.parent / "templates" / "distill.md"
    if bundled.is_file():
        return bundled.read_text(encoding="utf-8")

    return (
        "Distil the provided translation source material into a single "
        "Markdown article ready for the memory bank's inbox. Follow the "
        "frontmatter / structure conventions of the Supervertaler "
        "memory bank format."
    )


__all__ = ["DistillAgent", "DistillResult", "SourceFormat"]
